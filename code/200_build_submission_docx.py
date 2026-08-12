#!/usr/bin/env python3.11
# ============================================================
# Scientific question / 科学问题:
# 把 Markdown 文稿转成可直接投稿的 Word 文件(Nature 系:NEE / NC 通用格式)。
# Convert the Markdown manuscript into a submission-ready Word file in the
# Nature-family format shared by Nature Ecology & Evolution and Nature
# Communications.
#
# Objective / 分析目标:
#   1. 拆出题名页、正文、参考文献、表格、图注四块,按投稿顺序重排
#   2. pandoc 转 docx,保留上标引文与 Markdown 表格
#   3. python-docx 后处理:连续行号、双倍行距、Times New Roman 12 pt、
#      页码、表格独立分页 —— 这些都是 pandoc 做不到的 sectPr / 样式层设置
#   4. 输出字数核对表,与目标期刊的硬约束逐项比对
#
# Input / 输入:  analysis_v2/docs/MANUSCRIPT_v2.md
# Output / 输出: analysis_v2/submission/*.docx, *_wordcount.csv
#
# Key assumptions / 关键假设:
#   - Nature 系要求 Methods 置于参考文献之后,引文用编号上标 —— 与本文稿
#     现有结构一致,因此本脚本不改动章节顺序。
#   - 图件按 Nature 政策单独提交,正文只保留图注;表格作为真正的 Word 表格
#     嵌入,每张独立起页。
#   - 行号为连续编号(continuous),这是三家投稿系统的共同要求。
#
# Main packages: pandoc (外部), python-docx
# Run: python3.11 code/200_build_submission_docx.py
# ============================================================

from __future__ import annotations
import csv
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

V2 = Path("/Users/dingchenchen/Documents/New records/bird-new-distribution-records/"
          "tasks/bird_hazard_model_effort_upgrade_v2")
SRC = V2 / "analysis_v2/docs/MANUSCRIPT_v2.md"
OUT = V2 / "analysis_v2/submission"
OUT.mkdir(parents=True, exist_ok=True)

BODY_FONT = "Times New Roman"
BODY_PT = 12


def log(*a: object) -> None:
    print("[docx]", *a, flush=True)


# ------------------------------------------------------------
# 1. 拆分 Markdown
# ------------------------------------------------------------
def split_sections(text: str) -> dict[str, str]:
    """按二级标题切分文稿 / split the manuscript on level-2 headings."""
    heads = [m for m in re.finditer(r"^## (.+?)\s*$", text, re.M)]
    out: dict[str, str] = {}
    front = text[: heads[0].start()] if heads else text
    out["_front"] = front
    for i, m in enumerate(heads):
        end = heads[i + 1].start() if i + 1 < len(heads) else len(text)
        out[m.group(1).strip()] = text[m.end(): end].strip()
    return out


def word_count(t: str) -> int:
    """正文词数:剔除引文上标与表格行 / words excluding citations and table rows."""
    t = re.sub(r"<sup>.*?</sup>", "", t)
    t = re.sub(r"^\|.*\|\s*$", "", t, flags=re.M)
    t = re.sub(r"[*_`#>]", "", t)
    return len(re.findall(r"[A-Za-z0-9][A-Za-z0-9\-–—'’.]*", t))


# ------------------------------------------------------------
# 2. 组装投稿顺序的 Markdown
# ------------------------------------------------------------
def assemble(sec: dict[str, str]) -> tuple[str, str]:
    """返回 (正文 md, 表格与图注 md) / return body and back-matter markdown."""
    front = sec["_front"]
    m = re.search(r"^# (.+?)\s*$", front, re.M)
    title = m.group(1).strip() if m else "Untitled"
    authors = re.search(r"^\*\*(.+?)\*\*(.*)$", front, re.M)
    author_line = (authors.group(1) + authors.group(2)).strip() if authors else ""
    affil = re.search(r"^<sup>1</sup>\s*(.+?)$", front, re.M)
    affil_line = affil.group(1).strip() if affil else ""
    corr = re.search(r"^Correspondence:\s*(.+?)$", front, re.M)
    corr_line = corr.group(1).strip() if corr else ""

    body = [
        f"# {title}", "",
        author_line, "",
        f"<sup>1</sup> {affil_line}", "",
        f"Correspondence: {corr_line}", "",
        "## Abstract", "", sec.get("Abstract", ""), "",
        "## Introduction", "", sec.get("Introduction", ""), "",
        "## Results", "", sec.get("Results", ""), "",
        "## Discussion", "", sec.get("Discussion", ""), "",
        "## References", "", sec.get("References", ""), "",
        "## Methods", "", sec.get("Methods", ""), "",
    ]
    back = [
        "## Tables", "", sec.get("Tables", ""), "",
        "## Figure legends", "", sec.get("Figure legends", ""), "",
    ]
    return "\n".join(body), "\n".join(back)


# ------------------------------------------------------------
# 3. python-docx 后处理
# ------------------------------------------------------------
def add_line_numbers(section) -> None:
    """连续行号 / continuous line numbering — a sectPr setting pandoc cannot emit."""
    sect_pr = section._sectPr
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:start"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")
    sect_pr.append(ln)


def add_page_numbers(section) -> None:
    """页脚居中页码 / centred page number in the footer."""
    p = section.footer.paragraphs[0]
    p.alignment = 1
    for instr in ("begin", "instrText", "separate", "end"):
        el = OxmlElement(f"w:fld{instr}" if instr in ("begin", "separate", "end")
                         else "w:instrText")
        if instr == "instrText":
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el.set(qn("w:fldCharType"), instr)
            el.tag = qn("w:fldChar")
        run = p.add_run()
        run._r.append(el)


def style_document(path: Path) -> None:
    doc = Document(str(path))

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)

    for st in doc.styles:
        if st.name.startswith("Heading"):
            try:
                st.font.name = BODY_FONT
                st.font.size = Pt(BODY_PT + 1)
                st.font.bold = True
                st.font.color.rgb = None
            except AttributeError:
                pass

    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1)
        section.top_margin = section.bottom_margin = Inches(1)
        add_line_numbers(section)
        add_page_numbers(section)

    # 表格:统一字号并直接写入边框 / compact tables with explicit borders.
    # pandoc 生成的 docx 不含 "Table Grid" 样式,故直接写 tblBorders,
    # 不依赖任何预置样式名。
    def set_borders(tbl) -> None:
        tbl_pr = tbl._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:color"), "auto")
            borders.append(el)
        tbl_pr.append(borders)

    for tbl in doc.tables:
        set_borders(tbl)
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = BODY_FONT

    # 每张表前分页,但首张不加(否则文件第一页空白)
    # page break before each table caption except the first
    seen_table = False
    for para in doc.paragraphs:
        if re.match(r"^Table \d+\.", para.text.strip()):
            if seen_table:
                para.paragraph_format.page_break_before = True
            seen_table = True

    doc.save(str(path))


def to_pandoc_superscript(md: str) -> str:
    """把 <sup>…</sup> 改写成 pandoc 原生上标 ^…^。

    pandoc 只在输出 HTML 时透传 raw_html;转 docx 时会整块丢弃,
    引文上标会全部消失。必须先改写成原生语法。
    pandoc drops raw HTML when targeting docx, so citation superscripts
    vanish unless rewritten into its native ^…^ syntax first.
    """
    def repl(m: re.Match[str]) -> str:
        inner = m.group(1).strip()
        # 原生上标不允许未转义空格 / native superscript cannot contain bare spaces
        return "^" + inner.replace(" ", r"\ ") + "^"
    return re.sub(r"<sup>(.*?)</sup>", repl, md, flags=re.S)


def build(md: str, out_path: Path, title: str) -> None:
    md = to_pandoc_superscript(md)
    tmp = OUT / (out_path.stem + ".md")
    tmp.write_text(md, encoding="utf-8")
    cmd = ["pandoc", str(tmp), "-o", str(out_path),
           "--from", "markdown+raw_html+pipe_tables+tex_math_dollars",
           "--to", "docx", "--standalone"]
    r = subprocess.run(cmd, capture_output=True, text=True)
    if r.returncode != 0:
        sys.exit(f"pandoc 失败: {r.stderr[:500]}")
    style_document(out_path)
    log(f"{title} -> {out_path.name} ({out_path.stat().st_size/1024:.0f} KB)")


# ------------------------------------------------------------
# 4. 主流程
# ------------------------------------------------------------
def main() -> None:
    text = SRC.read_text(encoding="utf-8")
    sec = split_sections(text)
    body_md, back_md = assemble(sec)

    build(body_md, OUT / "Manuscript_NatureFamily_maintext.docx", "正文")
    build(back_md, OUT / "Manuscript_NatureFamily_tables_and_legends.docx", "表格与图注")

    # 字数核对 / word-count audit against journal limits
    counts = {k: word_count(sec.get(k, "")) for k in
              ("Abstract", "Introduction", "Results", "Discussion", "Methods")}
    main_txt = counts["Introduction"] + counts["Results"] + counts["Discussion"]
    n_ref = len(re.findall(r"^\d+\. ", sec.get("References", ""), re.M))
    n_tab = len(re.findall(r"^\*\*Table \d", sec.get("Tables", ""), re.M))
    n_fig = len(re.findall(r"^\*\*(Fig\.|Extended Data Fig\.)", sec.get("Figure legends", ""), re.M))

    rows = [
        ("Abstract words", counts["Abstract"], "~150", "≤150", "≤250"),
        ("Main text words (Intro+Results+Disc)", main_txt, "~3,500", "~5,000 (guide)", "—"),
        ("All-inclusive words (+Methods)", main_txt + counts["Methods"], "—", "—", "≤8,000"),
        ("References", n_ref, "~50 in main text", "no limit", "no limit"),
        ("Tables", n_tab, "", "", ""),
        ("Figure legends", n_fig, "", "", ""),
        ("Display items total", n_tab + n_fig, "≤6", "≤10", "reasonable"),
    ]
    csv_path = OUT / "wordcount_vs_journal_limits.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as fh:
        w = csv.writer(fh)
        w.writerow(["item", "current", "NEE", "NC", "GCB"])
        w.writerows(rows)

    log("字数核对:")
    for r in rows:
        print(f"  {r[0]:<38} {str(r[1]):>7}   NEE {r[2]:<16} NC {r[3]:<16} GCB {r[4]}")
    log("完成 / done ->", OUT)


if __name__ == "__main__":
    main()
