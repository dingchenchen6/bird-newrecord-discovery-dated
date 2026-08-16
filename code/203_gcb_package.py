#!/usr/bin/env python3.11
# ============================================================
# GCB 投稿全包装配器。
# Assemble the complete GCB submission package.
#
# 产出四件东西 / four deliverables:
#   1. figures/Figure1..Figure7 — 主图按 GCB 稿的编号重命名(PDF + PNG)
#      映射:GCB 1←源1, 2←源2, 3←源3, 4←源5, 5←源6, 6←源11(风玫瑰), 7←源M4
#   2. Manuscript_GCB_tables_and_legends.docx — 只含 Table 1–3 与 Figure 1–7 图注
#      (图注按 GCB 编号改写,引文转作者-年份)
#   3. Supporting_Information_GCB.docx — Notes S1–S6、Tables S1–S8、
#      Figures S1–S15(图注 + 嵌入 PNG)
#   4. data_code_availability.md — Data/Code availability 正式表述
#
# Note 的正文来自源文稿被移出的段落与模块文档,表格来自源稿 Tables 4–9
# 与结果 CSV;全部为既有文本的重组,不新写任何结论。
#
# Run: python3.11 code/203_gcb_package.py
# ============================================================

from __future__ import annotations
import importlib.util
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

V2 = Path("/Users/dingchenchen/Documents/New records/bird-new-distribution-records/"
          "tasks/bird_hazard_model_effort_upgrade_v2")
SRC = V2 / "analysis_v2/docs/MANUSCRIPT_v2.md"
A2 = V2 / "analysis_v2"
OUT = A2 / "submission/GCB"
FIGOUT = OUT / "figures"
FIGOUT.mkdir(parents=True, exist_ok=True)

BODY_FONT, BODY_PT = "Times New Roman", 12


def load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


cit = load_module("cit", V2 / "code/201_citation_tools.py")


def log(*a: object) -> None:
    print("[gcb]", *a, flush=True)


# ------------------------------------------------------------
# 0. 源文稿部件
# ------------------------------------------------------------
text = cit.normalise_a_labels(SRC.read_text(encoding="utf-8"))
REF_MAP = cit.build_ref_map(text)
heads = list(re.finditer(r"^## (.+?)\s*$", text, re.M))
PARTS: dict[str, str] = {}
for i, m in enumerate(heads):
    end = heads[i + 1].start() if i + 1 < len(heads) else len(text)
    PARTS[m.group(1).strip()] = text[m.end(): end].strip()


def grab_results_subsection(title_frag: str) -> str:
    """按小节标题片段取 Results 的整节文本。"""
    res = PARTS["Results"]
    m = re.search(rf"^### [^\n]*{re.escape(title_frag)}[^\n]*$", res, re.M)
    if not m:
        sys.exit(f"未找到 Results 小节: {title_frag}")
    nxt = re.search(r"^### ", res[m.end():], re.M)
    return res[m.end(): m.end() + nxt.start() if nxt else len(res)].strip()


def grab_table(n: int) -> str:
    """取源稿 Tables 段中的第 n 张表(标题 + 表体)。"""
    tb = PARTS["Tables"]
    m = re.search(rf"^\*\*Table {n}\.\*\*", tb, re.M)
    if not m:
        sys.exit(f"未找到 Table {n}")
    nxt = re.search(r"^\*\*Table \d+\.\*\*", tb[m.end():], re.M)
    return tb[m.start(): m.end() + nxt.start() if nxt else len(tb)].strip()


def grab_legend(tag: str) -> str:
    """取图注段中某一图的图注。tag 形如 'Fig. 1' / 'Fig. M4' / 'Extended Data Fig. 1'。"""
    fl = PARTS["Figure legends"]
    m = re.search(rf"^\*\*{re.escape(tag)} \|", fl, re.M)
    if not m:
        sys.exit(f"未找到图注: {tag}")
    nxt = re.search(r"^\*\*(?:Fig\.|Extended Data Fig\.)", fl[m.end():], re.M)
    return fl[m.start(): m.end() + nxt.start() if nxt else len(fl)].strip()


AD = lambda t: cit.to_author_date(t, REF_MAP)   # noqa: E731  作者-年份转换


# ------------------------------------------------------------
# 1. 主图重编号
# ------------------------------------------------------------
FIGMAP = [
    ("Figure1", A2 / "figures/Fig1_main_results_v2"),
    ("Figure2", A2 / "figures/Fig2_dating_correction_v2"),
    ("Figure3", A2 / "figures/Fig3_specification_robustness_v2"),
    ("Figure4", A2 / "figures/Fig5_model_anatomy_v2"),
    ("Figure5", A2 / "figures/Fig6_observation_process_v2"),
    ("Figure6", A2 / "figures_direction/overall/overall_direction_windrose"),
    ("Figure7", A2 / "figures_future/FigM4_mech_vs_ml_agreement_v2"),
]
for new, base in FIGMAP:
    for ext in ("pdf", "png"):
        srcf = base.with_suffix("." + ext)
        if srcf.exists():
            shutil.copy(srcf, FIGOUT / f"{new}.{ext}")
        else:
            log(f"警告: 缺 {srcf.name}")
log(f"主图 {len(FIGMAP)} 张已重编号 -> {FIGOUT}")

# ------------------------------------------------------------
# 2. 图注重编号(GCB 口径)
# ------------------------------------------------------------
LEGEND_SRC = ["Fig. 1", "Fig. 2", "Fig. 3", "Fig. 5", "Fig. 6", "Fig. 11", "Fig. M4"]
legends = []
for i, tag in enumerate(LEGEND_SRC, 1):
    body = grab_legend(tag)
    body = re.sub(rf"^\*\*{re.escape(tag)} \|", f"**Figure {i} |", body)
    # 图注内部对其他图的交叉引用一律按新编号改写
    body = (body.replace("Fig. 1", "Figure 1").replace("Fig. 2", "Figure 2")
                .replace("Fig. 3", "Figure 3").replace("Fig. 5", "Figure 4")
                .replace("Fig. 6", "Figure 5").replace("Fig. 11", "Figure 6")
                .replace("Fig. M4", "Figure 7"))
    legends.append(AD(body))

tables_md = "\n\n".join(AD(grab_table(n)) for n in (1, 2, 3))
back_md = "\n\n".join(["## Tables", tables_md, "## Figure legends"] + legends)

# ------------------------------------------------------------
# 3. Supporting Information
# ------------------------------------------------------------
def si_note(si_id: str, title: str, body_md: str) -> str:
    return f"### Supplementary {si_id}. {title}\n\n{body_md.strip()}"


module = lambda p: (V2 / p).read_text(encoding="utf-8")   # noqa: E731

NOTES = [
    si_note("Note S1", "Publication-lag stationarity and the reporting-completeness offset",
        grab_results_subsection("risk set and the dating problem") + "\n\n" +
        re.search(r"### Reporting completeness\s*(.*?)(?=^### )", PARTS["Methods"], re.M | re.S).group(1)),
    si_note("Note S2", "Migratory stratification",
        grab_results_subsection("Migratory strategy")),
    si_note("Note S3", "Species-level and province-level companion analyses",
        grab_results_subsection("Which species, and which provinces")),
    si_note("Note S4", "Resampling and null models: implementation detail",
        re.search(r"### Resampling and null models\s*(.*?)(?=^### )", PARTS["Methods"], re.M | re.S).group(1) +
        "\n\n" + grab_results_subsection("Model adequacy")),
    si_note("Note S5", "Robustness of the climate signal to window, baseline, source and indicator",
        grab_results_subsection("Robustness")),
    si_note("Note S6", "Records inside nature reserves: design, estimand ladder and limits",
        grab_results_subsection("nature reserves")),
]

# ---- SI 表:源稿 Tables 4–9 + 两张 CSV ----
def csv_to_md(path: Path, max_rows: int = 24) -> str:
    import csv as _csv
    rows = list(_csv.reader(path.open(encoding="utf-8")))
    if len(rows) > max_rows + 1:
        rows = rows[: max_rows + 1]
        trunc = True
    else:
        trunc = False
    out = ["| " + " | ".join(rows[0]) + " |",
           "|" + "---|" * len(rows[0])]
    out += ["| " + " | ".join(r) + " |" for r in rows[1:]]
    if trunc:
        out.append(f"\n*(前 {max_rows} 行;完整表见仓库 CSV)*")
    return "\n".join(out)


SI_TABLES = [
    ("Table S1", "Model adequacy checks", grab_table(4)),
    ("Table S2", "Ecological rationale for every component of the main model", grab_table(5)),
    ("Table S3", "Migratory stratification ladder and group-specific estimates", grab_table(6)),
    ("Table S4", "Species-level range measures under phylogenetic logistic regression", grab_table(7)),
    ("Table S5", "Province-level count and rate models", grab_table(8)),
    ("Table S6", "Species-level effects within taxonomic and migratory groups", grab_table(9)),
    ("Table S7", "Risk-set bridge from the released compilation to the modelling data",
     csv_to_md(A2 / "tables/tbl_riskset_bridge_v2.csv")),
    ("Table S8", "Fine grid of accumulation windows (3–23 years), annual mean temperature",
     csv_to_md(A2 / "tables/_iw_tavg_annual.csv")),
]
si_tables_md = []
for sid, title, body in SI_TABLES:
    body = re.sub(r"^\*\*Table \d+\.\*\*\s*", "", body)          # 去旧编号
    si_tables_md.append(f"### Supplementary {sid}. {title}\n\n{AD(body)}")

# ---- SI 图:图注 + 嵌图 ----
SI_FIGS = [
    ("Fig. S1", A2 / "figures/Fig4_random_structure_v2.png", "Fig. 4"),
    ("Fig. S2", A2 / "figures/Fig7_window_baseline_sensitivity_v2.png", "Fig. 7"),
    ("Fig. S3", A2 / "figures/Fig8_migratory_strategy_v2.png", "Fig. 8"),
    ("Fig. S4", A2 / "figures/Fig9_species_level_v2.png", "Fig. 9"),
    ("Fig. S5", A2 / "figures/Fig10_province_level_v2.png", "Fig. 10"),
    ("Fig. S6", A2 / "figures_future/FigM1_future_mechanistic_v2.png", "Fig. M1"),
    ("Fig. S7", A2 / "figures_future/FigM2_future_ml_v2.png", "Fig. M2"),
    ("Fig. S8", A2 / "figures_future/FigM3_shap_interpretability_v2.png", "Fig. M3"),
    ("Fig. S9", A2 / "figures/FigS1_dharma_v2.png", "Extended Data Fig. 1"),
    ("Fig. S10", A2 / "figures_future/FigS3_future_unmasked_v2.png", "Extended Data Fig. 3"),
    ("Fig. S11", A2 / "figures/FigS5_resampling_v2.png", "Extended Data Fig. 5"),
    ("Fig. S12", A2 / "figures_pa/FigP1_pa_enrichment.png", None),
    ("Fig. S13", A2 / "figures_pa/FigP2_pa_persistence.png", None),
    ("Fig. S14", A2 / "figures_pa/FigP3_pa_mismatch_monitoring.png", None),
    ("Fig. S15", V2 / "analysis_v2/figures_direction/combined", None),
]
PA_LEGENDS = {
    "Fig. S12": "**Fig. S12 | Enrichment of new records inside mapped nature reserves, decomposed against the observation process.** **a**, Share of records, checklist-weighted birding effort, birding locations and land area falling inside the 1,028 mapped reserves. **b**, Conditional logistic estimand ladder matching each event to 200 same-province same-year checklist-weighted control locations. **c**, Robustness: host-reserve exclusions, period splits and discovery-method subsets. **d**, Concentration of inside-reserve events among host reserves.",
    "Fig. S13": "**Fig. S13 | Reserve status does not raise the probability that a first record becomes a persistent presence.** Re-detection of the same species in the same province in the GBIF/eBird layer, raw and adjusted for subsequent survey effort, with the pre-declared equivalence bounds.",
    "Fig. S14": "**Fig. S14 | The mapped reserve network shows no measurable mismatch with observed cell-level warming, and the monitoring gap decomposes into graftable and non-graftable areas.** **a**, Area-weighted reserve coverage across warming quintiles under the provincial-broadcast and the recomputed cell-level warming fields. **b**, Stratified permutation test within province × elevation bands. **c**–**d**, Monitoring-grafting priority classes.",
    "Fig. S15": "**Fig. S15 | Per-order directionality of new records relative to species' range centroids.** Wind-rose panels for the sixteen orders, each carrying its number of records and species.",
}


def find_s15() -> Path | None:
    comb = A2 / "figures_direction/combined"
    for cand in sorted(comb.glob("*.png")):
        if "order" in cand.name or "facet" in cand.name or "wind" in cand.name:
            return cand
    pngs = sorted(comb.glob("*.png"))
    return pngs[0] if pngs else None


si_figs_md = []
for sid, path, src_tag in SI_FIGS:
    if sid == "Fig. S15":
        path = find_s15()
        if path is None:
            log("警告: 未找到 S15 组合图"); continue
    if not Path(path).exists():
        log(f"警告: 缺 {path}"); continue
    if src_tag:
        body = grab_legend(src_tag)
        body = re.sub(r"^\*\*(?:Fig\.|Extended Data Fig\.) [^|]*\|", f"**{sid} |", body)
    else:
        body = PA_LEGENDS[sid]
    si_figs_md.append(f"![]({Path(path).as_posix()})\n\n{AD(body)}")

def remap_si_refs(md: str) -> str:
    """把 SI 内文里源稿编号的交叉引用改写成 GCB 编号。

    次序敏感:先替换多位数与带前缀的标签,再替换单位数,避免 'Fig. 1' 抢先
    匹配 'Fig. 10'。Rewrite source-numbered cross-references to GCB numbering;
    longer tags first so 'Fig. 10' is not eaten by 'Fig. 1'.
    """
    import re as _re
    RULES = [
        ("Extended Data Table 2", "Table S8"),
        ("Extended Data Table 1", "Table S7"),
        ("Extended Data Fig. 5", "Fig. S11"),
        ("Extended Data Fig. 3", "Fig. S10"),
        ("Extended Data Fig. 1", "Fig. S9"),
        ("Fig. M1", "Fig. S6"), ("Fig. M2", "Fig. S7"),
        ("Fig. M3", "Fig. S8"), ("Fig. M4", "Figure 7"),
        ("Fig. P1", "Fig. S12"), ("Fig. P2", "Fig. S13"), ("Fig. P3", "Fig. S14"),
        ("Fig. 11", "Figure 6"), ("Fig. 10", "Fig. S5"),
        ("Fig. 9", "Fig. S4"), ("Fig. 8", "Fig. S3"), ("Fig. 7", "Fig. S2"),
        ("Fig. 6", "Figure 5"), ("Fig. 5", "Figure 4"), ("Fig. 4", "Fig. S1"),
        ("Fig. 3", "Figure 3"), ("Fig. 2", "Figure 2"), ("Fig. 1", "Figure 1"),
        ("Table 9", "Table S6"), ("Table 8", "Table S5"), ("Table 7", "Table S4"),
        ("Table 6", "Table S3"), ("Table 5", "Table S2"), ("Table 4", "Table S1"),
        # 主表 1–3 编号未变,但 SI 中引用应指明是正文表
        ("Table 2", "Table 2 (main text)"), ("Table 1", "Table 1 (main text)"),
        ("Table 3", "Table 3 (main text)"),
    ]
    for old_tag, new_tag in RULES:
        # 只替换后面不紧跟数字的位置(防 Fig. 1 吃掉 Fig. 15 之类)
        md = _re.sub(_re.escape(old_tag) + r"(?!\d)", new_tag, md)
    return md


si_md = "\n\n".join([
    "# Supporting Information",
    "",
    "**Warming and survey effort contribute comparably to the generation of new bird "
    "distribution records once records are dated to discovery**",
    "",
    "Chenchen Ding, et al.",
    "",
    "## Supplementary Notes",
    AD("\n\n".join(NOTES)),
    "## Supplementary Tables",
    "\n\n".join(si_tables_md),
    "## Supplementary Figures",
    "\n\n".join(si_figs_md),
])

# ------------------------------------------------------------
# 4. Data / Code availability
# ------------------------------------------------------------
avail = """## Data availability

The released compilation of provincial new bird records (CBNR), the harmonised
species-trait table, the effort panel, the derived climate panels and every result
table reported in this paper (70+ CSV files, one per reported number) are available
at https://github.com/dingchenchen6/bird-newrecord-discovery-dated. BirdLife
International range polygons and the Chinese nature-reserve boundary layer are
redistributed under their providers' licences and are referenced by version in the
repository README. CRU TS 4.09, WorldClim 2.1 and CMIP6 (WorldClim downscaling)
climate surfaces are publicly available from their providers.

## Code availability

The full analysis pipeline (R scripts 130–203, run in order) reproduces every
number, table and figure in this paper from the released inputs, including a smoke
test that re-derives the four main-model coefficients. Archived at the same
repository.
"""
(OUT / "data_code_availability.md").write_text(avail, encoding="utf-8")

si_md = remap_si_refs(si_md)

# ------------------------------------------------------------
# 5. 输出 Word
# ------------------------------------------------------------
def style_document(path: Path, landscape_tables: bool = False) -> None:
    doc = Document(str(path))
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    for st in doc.styles:
        if st.name.startswith("Heading"):
            try:
                st.font.name, st.font.size, st.font.bold = BODY_FONT, Pt(BODY_PT + 1), True
                st.font.color.rgb = RGBColor(0, 0, 0)
            except AttributeError:
                pass
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1)
        section.top_margin = section.bottom_margin = Inches(1)
        p = section.footer.paragraphs[0]
        p.alignment = 1
        for kind in ("begin", "instrText", "separate", "end"):
            el = OxmlElement("w:instrText" if kind == "instrText" else "w:fldChar")
            if kind == "instrText":
                el.set(qn("xml:space"), "preserve"); el.text = " PAGE "
            else:
                el.set(qn("w:fldCharType"), kind)
            p.add_run()._r.append(el)
    for tbl in doc.tables:
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4"); e.set(qn("w:color"), "auto")
            borders.append(e)
        tbl._tbl.tblPr.append(borders)
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.size, run.font.name = Pt(8), BODY_FONT
    doc.save(str(path))


def build(md: str, out_path: Path) -> None:
    md = re.sub(r"<sup>(.*?)</sup>",
                lambda m: "^" + m.group(1).strip().replace(" ", r"\ ") + "^", md, flags=re.S)
    tmp = out_path.with_suffix(".md")
    tmp.write_text(md, encoding="utf-8")
    r = subprocess.run(["pandoc", str(tmp), "-o", str(out_path),
                        "--from", "markdown+raw_html+pipe_tables+implicit_figures",
                        "--to", "docx", "--standalone",
                        "--resource-path", str(V2)],
                       capture_output=True, text=True)
    if r.returncode != 0:
        sys.exit(f"pandoc 失败 ({out_path.name}): {r.stderr[:400]}")
    style_document(out_path)
    log(f"{out_path.name} ({out_path.stat().st_size/1024:.0f} KB)")


build(back_md, OUT / "Manuscript_GCB_tables_and_legends.docx")
build(si_md, OUT / "Supporting_Information_GCB.docx")
log("完成 ->", OUT)
