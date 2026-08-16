#!/usr/bin/env python3.11
# ============================================================
# NEE 与 NC 投稿包装配器(与 203 的 GCB 包同构,Nature 系体例)。
# Assemble the NEE and NC submission packages, Nature-family style.
#
# 每版产出 / per journal:
#   figures/FigN.(pdf,png)                 主图按该版编号重命名
#   Manuscript_{J}_tables_and_legends.docx 正文表 + 主图图注(该版编号)
#   Supplementary_Information_{J}.docx     Notes + Tables + Figures(嵌图)
#
# 编号映射由各版 revised_*.json 的 si_manifest 与正文引用共同决定;
# 交叉引用重写采用两阶段替换(源标签 → 占位符 → 新标签),
# 避免 Fig.1 ↔ Fig.2 之类互换时的串号。
# Cross-references are rewritten in two phases (source tag → token → new tag)
# so that swapped numbers cannot collide.
#
# Run: python3.11 code/204_nee_nc_packages.py
# ============================================================

from __future__ import annotations
import importlib.util
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

V2 = Path("/Users/dingchenchen/Documents/New records/bird-new-distribution-records/"
          "tasks/bird_hazard_model_effort_upgrade_v2")
A2 = V2 / "analysis_v2"
SRC = A2 / "docs/MANUSCRIPT_v2.md"
BODY_FONT, BODY_PT = "Times New Roman", 12


def load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


cit = load_module("cit", V2 / "code/201_citation_tools.py")


def log(*a: object) -> None:
    print("[nee-nc]", *a, flush=True)


# ---------------- 源文稿部件 ----------------
text = cit.normalise_a_labels(SRC.read_text(encoding="utf-8"))
heads = list(re.finditer(r"^## (.+?)\s*$", text, re.M))
PARTS: dict[str, str] = {}
for i, m in enumerate(heads):
    end = heads[i + 1].start() if i + 1 < len(heads) else len(text)
    PARTS[m.group(1).strip()] = text[m.end(): end].strip()


def results_sub(frag: str) -> str:
    res = PARTS["Results"]
    m = re.search(rf"^### [^\n]*{re.escape(frag)}[^\n]*$", res, re.M)
    if not m:
        sys.exit(f"未找到小节 {frag}")
    nxt = re.search(r"^### ", res[m.end():], re.M)
    return res[m.end(): m.end() + nxt.start() if nxt else len(res)].strip()


def methods_sub(frag: str) -> str:
    m = re.search(rf"### {re.escape(frag)}\s*(.*?)(?=^### |\Z)", PARTS["Methods"], re.M | re.S)
    if not m:
        sys.exit(f"未找到 Methods 小节 {frag}")
    return m.group(1).strip()


def src_table(n: int) -> str:
    tb = PARTS["Tables"]
    m = re.search(rf"^\*\*Table {n}\.\*\*", tb, re.M)
    nxt = re.search(r"^\*\*Table \d+\.\*\*", tb[m.end():], re.M)
    body = tb[m.start(): m.end() + nxt.start() if nxt else len(tb)].strip()
    return re.sub(r"^\*\*Table \d+\.\*\*\s*", "", body)


def src_legend(tag: str) -> str:
    fl = PARTS["Figure legends"]
    m = re.search(rf"^\*\*{re.escape(tag)} \|", fl, re.M)
    if not m:
        sys.exit(f"未找到图注 {tag}")
    nxt = re.search(r"^\*\*(?:Fig\.|Extended Data Fig\.)", fl[m.end():], re.M)
    return fl[m.start(): m.end() + nxt.start() if nxt else len(fl)].strip()


def csv_to_md(path: Path, max_rows: int = 24) -> str:
    import csv as _csv
    rows = list(_csv.reader(path.open(encoding="utf-8")))
    trunc = len(rows) > max_rows + 1
    rows = rows[: max_rows + 1]
    out = ["| " + " | ".join(rows[0]) + " |", "|" + "---|" * len(rows[0])]
    out += ["| " + " | ".join(r) + " |" for r in rows[1:]]
    if trunc:
        out.append(f"\n*(first {max_rows} rows; full table in the repository CSV)*")
    return "\n".join(out)


def two_phase_remap(md: str, rules: list[tuple[str, str]]) -> str:
    """两阶段替换:先全部替换成唯一占位符,再落到目标标签。"""
    tokens = []
    for i, (old, new) in enumerate(rules):
        tok = f"\x00REF{i}\x00"
        md = re.sub(re.escape(old) + r"(?!\d)", tok, md)
        tokens.append((tok, new))
    for tok, new in tokens:
        md = md.replace(tok, new)
    return md


# ---------------- 图件源 ----------------
F = {
    "main1": A2 / "figures/Fig1_main_results_v2",
    "dating": A2 / "figures/Fig2_dating_correction_v2",
    "robust": A2 / "figures/Fig3_specification_robustness_v2",
    "re": A2 / "figures/Fig4_random_structure_v2",
    "anatomy": A2 / "figures/Fig5_model_anatomy_v2",
    "obs": A2 / "figures/Fig6_observation_process_v2",
    "window": A2 / "figures/Fig7_window_baseline_sensitivity_v2",
    "migr": A2 / "figures/Fig8_migratory_strategy_v2",
    "species": A2 / "figures/Fig9_species_level_v2",
    "prov": A2 / "figures/Fig10_province_level_v2",
    "wind": A2 / "figures_direction/overall/overall_direction_windrose",
    "windfac": A2 / "figures_direction/combined/order_direction_windrose_facets",
    "m1": A2 / "figures_future/FigM1_future_mechanistic_v2",
    "m2": A2 / "figures_future/FigM2_future_ml_v2",
    "m3": A2 / "figures_future/FigM3_shap_interpretability_v2",
    "m4": A2 / "figures_future/FigM4_mech_vs_ml_agreement_v2",
    "dharma": A2 / "figures/FigS1_dharma_v2",
    "unmask": A2 / "figures_future/FigS3_future_unmasked_v2",
    "resamp": A2 / "figures/FigS5_resampling_v2",
    "pa1": A2 / "figures_pa/FigP1_pa_enrichment",
    "pa2": A2 / "figures_pa/FigP2_pa_persistence",
    "pa3": A2 / "figures_pa/FigP3_pa_mismatch_monitoring",
}
PA_LEG = {
    "pa1": "**{sid} | Enrichment of new records inside mapped nature reserves, decomposed against the observation process.** **a**, Share of records, checklist-weighted birding effort, birding locations and land area inside the 1,028 mapped reserves. **b**, Conditional logistic estimand ladder, each event matched to 200 same-province same-year checklist-weighted control locations. **c**, Robustness across host-reserve exclusions, period splits and discovery-method subsets. **d**, Concentration of inside-reserve events among host reserves.",
    "pa2": "**{sid} | Reserve status does not raise the probability that a first record becomes a persistent presence.** Re-detection of the same species in the same province in the GBIF/eBird layer, raw and adjusted for subsequent survey effort, with pre-declared equivalence bounds.",
    "pa3": "**{sid} | The mapped reserve network shows no measurable mismatch with observed cell-level warming.** **a**, Area-weighted reserve coverage across warming quintiles under provincial-broadcast versus recomputed cell-level warming. **b**, Stratified permutation within province × elevation bands. **c**–**d**, Monitoring-grafting priority classes.",
    "windfac": "**{sid} | Per-order directionality of new records relative to species' range centroids.** Wind-rose panels for the sixteen best-sampled orders, each annotated with its number of records and species.",
}

# ---------------- 各版配置 ----------------
# 主图:(新编号名, 源图键, 源图注标签)
# SI 图:(SI 编号, 源图键, 源图注标签或 None=用 PA_LEG)
CONFIG = {
    "NEE": {
        "main_figs": [("Fig1", "main1", "Fig. 1"), ("Fig2", "dating", "Fig. 2"),
                      ("Fig3", "obs", "Fig. 6"), ("Fig4", "wind", "Fig. 11"),
                      ("Fig5", "m4", "Fig. M4")],
        "main_tables": [(1, 2)],           # (新编号, 源编号)
        "si_figs": [("Fig. S1", "re", "Fig. 4"), ("Fig. S2", "robust", "Fig. 3"),
                    ("Fig. S3", "window", "Fig. 7"), ("Fig. S4", "anatomy", "Fig. 5"),
                    ("Fig. S5", "migr", "Fig. 8"), ("Fig. S6", "species", "Fig. 9"),
                    ("Fig. S7", "prov", "Fig. 10"), ("Fig. S8", "dharma", "Extended Data Fig. 1"),
                    ("Fig. S9", "resamp", "Extended Data Fig. 5"), ("Fig. S10", "m1", "Fig. M1"),
                    ("Fig. S11", "unmask", "Extended Data Fig. 3"), ("Fig. S12", "m2", "Fig. M2"),
                    ("Fig. S13", "pa1", None), ("Fig. S14", "pa2", None), ("Fig. S15", "pa3", None),
                    ("Fig. S16", "windfac", None), ("Fig. S17", "m3", "Fig. M3")],
        "si_tables": [("Table S1", "Step-by-step decomposition S0–S5", src_table(1)),
                      ("Table S2", "Random-effect structure comparison on four criteria", src_table(3)),
                      ("Table S3", "Model adequacy checks", src_table(4)),
                      ("Table S4", "Ecological rationale for every model component", src_table(5)),
                      ("Table S5", "Migratory stratification", src_table(6)),
                      ("Table S6", "Species-level range measures", src_table(7)),
                      ("Table S7", "Province-level models", src_table(8)),
                      ("Table S8", "Species-level effects within groups", src_table(9)),
                      ("Table S9", "Risk-set bridge from compilation to modelling data",
                       csv_to_md(A2 / "tables/tbl_riskset_bridge_v2.csv")),
                      ("Table S10", "Fine accumulation-window grid, 3–23 years",
                       csv_to_md(A2 / "tables/_iw_tavg_annual.csv"))],
        "si_notes": [
            ("Note S1", "Risk-set construction, publication-lag stationarity and the completeness offset",
             results_sub("risk set and the dating problem") + "\n\n" + methods_sub("Reporting completeness") + "\n\n" + methods_sub("Risk set")),
            ("Note S2", "Importance criteria and the species-specific climate gradient",
             results_sub("Warming and effort contribute comparably")),
            ("Note S3", "Random-effect structure selection",
             results_sub("Separating the observation process") + "\n\n" + results_sub("What each level of the model means")),
            ("Note S4", "Migratory stratification", results_sub("Migratory strategy")),
            ("Note S5", "Species-level and province-level companion analyses",
             results_sub("Which species, and which provinces")),
            ("Note S6", "Per-order tests of directional uniformity",
             results_sub("east and northeast")),
            ("Note S7", "Full robustness suite", results_sub("Robustness")),
            ("Note S8", "Model adequacy and resampling",
             results_sub("Model adequacy") + "\n\n" + methods_sub("Resampling and null models")),
            ("Note S9", "Projection detail and covariate support",
             results_sub("How far can these relationships be projected") + "\n\n" + methods_sub("Projections")),
            ("Note S10", "Nature-reserve conditional logistic analysis",
             results_sub("nature reserves")),
        ],
        # 源标签 → NEE 标签(两阶段)
        "remap": [
            ("Extended Data Table 2", "Supplementary Table S10"),
            ("Extended Data Table 1", "Supplementary Table S9"),
            ("Extended Data Fig. 5", "Supplementary Fig. S9"),
            ("Extended Data Fig. 3", "Supplementary Fig. S11"),
            ("Extended Data Fig. 1", "Supplementary Fig. S8"),
            ("Fig. M1", "Supplementary Fig. S10"), ("Fig. M2", "Supplementary Fig. S12"),
            ("Fig. M3", "Supplementary Fig. S17"), ("Fig. M4", "Fig. 5"),
            ("Fig. P1", "Supplementary Fig. S13"), ("Fig. P2", "Supplementary Fig. S14"),
            ("Fig. P3", "Supplementary Fig. S15"),
            ("Fig. 11", "Fig. 4"), ("Fig. 10", "Supplementary Fig. S7"),
            ("Fig. 9", "Supplementary Fig. S6"), ("Fig. 8", "Supplementary Fig. S5"),
            ("Fig. 7", "Supplementary Fig. S3"), ("Fig. 6", "Fig. 3"),
            ("Fig. 5", "Supplementary Fig. S4"), ("Fig. 4", "Supplementary Fig. S1"),
            ("Fig. 3", "Supplementary Fig. S2"), ("Fig. 2", "Fig. 2"), ("Fig. 1", "Fig. 1"),
            ("Table 9", "Supplementary Table S8"), ("Table 8", "Supplementary Table S7"),
            ("Table 7", "Supplementary Table S6"), ("Table 6", "Supplementary Table S5"),
            ("Table 5", "Supplementary Table S4"), ("Table 4", "Supplementary Table S3"),
            ("Table 3", "Supplementary Table S2"), ("Table 2", "Table 1"),
            ("Table 1", "Supplementary Table S1"),
        ],
    },
    "NC": {
        "main_figs": [("Fig1", "dating", "Fig. 2"), ("Fig2", "main1", "Fig. 1"),
                      ("Fig3", "anatomy", "Fig. 5"), ("Fig4", "obs", "Fig. 6"),
                      ("Fig5", "wind", "Fig. 11"), ("Fig6", "robust", "Fig. 3"),
                      ("Fig7", "m4", "Fig. M4")],
        "main_tables": [(1, 1), (2, 2), (3, 3)],
        "si_figs": [("Fig. S1", "re", "Fig. 4"), ("Fig. S2", "window", "Fig. 7"),
                    ("Fig. S3", "migr", "Fig. 8"), ("Fig. S4", "species", "Fig. 9"),
                    ("Fig. S5", "prov", "Fig. 10"), ("Fig. S6", "m1", "Fig. M1"),
                    ("Fig. S7", "m2", "Fig. M2"), ("Fig. S8", "m3", "Fig. M3"),
                    ("Fig. S9", "dharma", "Extended Data Fig. 1"),
                    ("Fig. S10", "unmask", "Extended Data Fig. 3"),
                    ("Fig. S11", "resamp", "Extended Data Fig. 5"),
                    ("Fig. S12", "pa1", None), ("Fig. S13", "pa2", None), ("Fig. S14", "pa3", None)],
        "si_tables": [("Table S1", "Model adequacy checks", src_table(4)),
                      ("Table S2", "Ecological rationale for every model component", src_table(5)),
                      ("Table S3", "Migratory stratification", src_table(6)),
                      ("Table S4", "Species-level range measures", src_table(7)),
                      ("Table S5", "Province-level models", src_table(8)),
                      ("Table S6", "Species-level effects within groups", src_table(9)),
                      ("Table S7", "Risk-set bridge from compilation to modelling data",
                       csv_to_md(A2 / "tables/tbl_riskset_bridge_v2.csv")),
                      ("Table S8", "Fine accumulation-window grid, 3–23 years",
                       csv_to_md(A2 / "tables/_iw_tavg_annual.csv"))],
        "si_notes": [
            ("Note S1", "Publication-lag stationarity and the completeness offset",
             results_sub("risk set and the dating problem") + "\n\n" + methods_sub("Reporting completeness")),
            ("Note S2", "Nature reserves: matched conditional-logistic design",
             results_sub("nature reserves")),
            ("Note S3", "Migratory stratification", results_sub("Migratory strategy")),
            ("Note S4", "Species-level and province-level companion analyses",
             results_sub("Which species, and which provinces")),
            ("Note S5", "Robustness detail", results_sub("Robustness")),
            ("Note S6", "Permutation nulls, bootstraps and model adequacy",
             methods_sub("Resampling and null models") + "\n\n" + results_sub("Model adequacy")),
            ("Note S7", "Projection design and covariate support",
             results_sub("How far can these relationships be projected") + "\n\n" + methods_sub("Projections")),
        ],
        "remap": [
            ("Extended Data Table 2", "Supplementary Table S8"),
            ("Extended Data Table 1", "Supplementary Table S7"),
            ("Extended Data Fig. 5", "Supplementary Fig. S11"),
            ("Extended Data Fig. 3", "Supplementary Fig. S10"),
            ("Extended Data Fig. 1", "Supplementary Fig. S9"),
            ("Fig. M1", "Supplementary Fig. S6"), ("Fig. M2", "Supplementary Fig. S7"),
            ("Fig. M3", "Supplementary Fig. S8"), ("Fig. M4", "Fig. 7"),
            ("Fig. P1", "Supplementary Fig. S12"), ("Fig. P2", "Supplementary Fig. S13"),
            ("Fig. P3", "Supplementary Fig. S14"),
            ("Fig. 11", "Fig. 5"), ("Fig. 10", "Supplementary Fig. S5"),
            ("Fig. 9", "Supplementary Fig. S4"), ("Fig. 8", "Supplementary Fig. S3"),
            ("Fig. 7", "Supplementary Fig. S2"), ("Fig. 6", "Fig. 4"),
            ("Fig. 5", "Fig. 3"), ("Fig. 4", "Supplementary Fig. S1"),
            ("Fig. 3", "Fig. 6"), ("Fig. 2", "Fig. 1"), ("Fig. 1", "Fig. 2"),
            ("Table 9", "Supplementary Table S6"), ("Table 8", "Supplementary Table S5"),
            ("Table 7", "Supplementary Table S4"), ("Table 6", "Supplementary Table S3"),
            ("Table 5", "Supplementary Table S2"), ("Table 4", "Supplementary Table S1"),
        ],
    },
}


# ---------------- Word 输出 ----------------
def style_document(path: Path) -> None:
    doc = Document(str(path))
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    for st in doc.styles:
        if st.name.startswith("Heading"):
            try:
                st.font.name, st.font.size, st.font.bold = BODY_FONT, Pt(BODY_PT + 1), True
                st.font.color.rgb = RGBColor(0, 0, 0)
            except AttributeError:
                pass
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1)
        section.top_margin = section.bottom_margin = Inches(1)
        p = section.footer.paragraphs[0]
        p.alignment = 1
        for kind in ("begin", "instrText", "separate", "end"):
            el = OxmlElement("w:instrText" if kind == "instrText" else "w:fldChar")
            if kind == "instrText":
                el.set(qn("xml:space"), "preserve"); el.text = " PAGE "
            else:
                el.set(qn("w:fldCharType"), kind)
            p.add_run()._r.append(el)
    for tbl in doc.tables:
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4"); e.set(qn("w:color"), "auto")
            borders.append(e)
        tbl._tbl.tblPr.append(borders)
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.size, run.font.name = Pt(8), BODY_FONT
    doc.save(str(path))


def build(md: str, out_path: Path) -> None:
    md = re.sub(r"<sup>(.*?)</sup>",
                lambda m: "^" + m.group(1).strip().replace(" ", r"\ ") + "^", md, flags=re.S)
    tmp = out_path.with_suffix(".md")
    tmp.write_text(md, encoding="utf-8")
    r = subprocess.run(["pandoc", str(tmp), "-o", str(out_path),
                        "--from", "markdown+raw_html+pipe_tables+implicit_figures",
                        "--to", "docx", "--standalone", "--resource-path", str(V2)],
                       capture_output=True, text=True)
    if r.returncode != 0:
        sys.exit(f"pandoc 失败 ({out_path.name}): {r.stderr[:400]}")
    style_document(out_path)
    log(f"{out_path.name} ({out_path.stat().st_size/1024:.0f} KB)")


# ---------------- 主流程 ----------------
for J, cfg in CONFIG.items():
    OUT = A2 / f"submission/{J}"
    FIGOUT = OUT / "figures"
    FIGOUT.mkdir(parents=True, exist_ok=True)
    remap = lambda md: two_phase_remap(md, cfg["remap"])   # noqa: E731

    # 1) 主图
    for newname, key, _ in cfg["main_figs"]:
        for ext in ("pdf", "png"):
            f = F[key].with_suffix("." + ext)
            if f.exists():
                shutil.copy(f, FIGOUT / f"{newname}.{ext}")
            else:
                log(f"警告: 缺 {f.name}")
    log(f"{J}: 主图 {len(cfg['main_figs'])} 张")

    # 2) 表格与图注
    tables_md = []
    for newn, srcn in cfg["main_tables"]:
        body = remap(src_table(srcn))
        tables_md.append(f"**Table {newn}.** {body}")
    legends_md = []
    for i, (_, key, legtag) in enumerate(cfg["main_figs"], 1):
        # 必须先在源编号语境下做 remap,再设置新标题;
        # 否则新标题"Fig. 3"会被映射规则二次改写,编号全部错乱。
        # Remap in the source-numbering context FIRST, then stamp the new
        # heading — the other order sends the fresh heading through the map.
        body = remap(src_legend(legtag)) if legtag else PA_LEG[key].format(sid=f"Fig. {i}")
        body = re.sub(r"^\*\*(?:Fig\.|Extended Data Fig\.|Supplementary Fig\.) [^|]*\|",
                      f"**Fig. {i} |", body)
        legends_md.append(body)
    back_md = "\n\n".join(["## Tables"] + tables_md + ["## Figure legends"] + legends_md)
    build(back_md, OUT / f"Manuscript_{J}_tables_and_legends.docx")

    # 3) SI 文档
    notes_md = [f"### Supplementary {nid}. {title}\n\n{remap(body)}"
                for nid, title, body in cfg["si_notes"]]
    si_tables_md = [f"### Supplementary {tid}. {title}\n\n{remap(body)}"
                    for tid, title, body in cfg["si_tables"]]
    si_figs_md = []
    for sid, key, legtag in cfg["si_figs"]:
        png = F[key].with_suffix(".png")
        if not png.exists():
            log(f"警告: 缺 {png.name}"); continue
        if legtag:
            body = remap(src_legend(legtag))          # 同理:先 remap 后设标题
            body = re.sub(r"^\*\*(?:Fig\.|Extended Data Fig\.|Supplementary Fig\.) [^|]*\|",
                          f"**{sid} |", body)
        else:
            body = PA_LEG[key].format(sid=sid)
        si_figs_md.append(f"![]({png.as_posix()})\n\n{body}")
    si_md = "\n\n".join([
        "# Supplementary Information", "",
        "**Warming and survey effort contribute comparably to the generation of new bird "
        "distribution records once records are dated to discovery**", "",
        "Chenchen Ding, et al.", "",
        "## Supplementary Notes", "\n\n".join(notes_md),
        "## Supplementary Tables", "\n\n".join(si_tables_md),
        "## Supplementary Figures", "\n\n".join(si_figs_md),
    ])
    build(si_md, OUT / f"Supplementary_Information_{J}.docx")

    # 4) 残留旧编号扫描(排除映射后合法的目标标签)
    leftovers = re.findall(r"(?:Extended Data (?:Fig\.|Table)|Fig\. [MP]\d)", si_md)
    log(f"{J}: SI 残留旧标签 {len(leftovers)} 处" + (f" {set(leftovers)}" if leftovers else ""))

log("完成")
