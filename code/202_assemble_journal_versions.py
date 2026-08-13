#!/usr/bin/env python3.11
# ============================================================
# 把三家期刊各自的改写稿装配成完整 Word 文件。
# Assemble the three journal-specific drafts into complete Word files.
#
# NEE / NC:  题名页 → Abstract → Introduction → Results → Discussion
#            → References → Methods,编号上标引文(Nature 系顺序)
# GCB:       题名页 → Abstract → Introduction → Materials and Methods
#            → Results → Discussion → References,作者-年份引文、文献按姓氏排序
#
# 引文体例转换用确定性代码(201_citation_tools),不交给语言模型,
# 且只作用于正文各段,绝不触碰题名页的作者单位上标 <sup>1</sup>。
#
# Input:  analysis_v2/submission/work/revised_{NEE,NC,GCB}.json
#         analysis_v2/submission/work/plan_{NEE,NC,GCB}.json   (摘要)
#         analysis_v2/docs/MANUSCRIPT_v2.md                    (Methods/References/表/图注)
# Output: analysis_v2/submission/{NEE,NC,GCB}/*.docx + wordcount.csv
#
# Run: python3.11 code/202_assemble_journal_versions.py
# ============================================================

from __future__ import annotations
import csv
import importlib.util
import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

V2 = Path("/Users/dingchenchen/Documents/New records/bird-new-distribution-records/"
          "tasks/bird_hazard_model_effort_upgrade_v2")
SRC = V2 / "analysis_v2/docs/MANUSCRIPT_v2.md"
SUB = V2 / "analysis_v2/submission"
WORK = SUB / "work"
CODE = V2 / "code"

BODY_FONT, BODY_PT = "Times New Roman", 12
JOURNALS = {
    "NEE": ("Nature Ecology & Evolution", 3600),
    "NC": ("Nature Communications", 5000),
    "GCB": ("Global Change Biology", 8000),
}


def load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


cit = load_module("cit", CODE / "201_citation_tools.py")


def log(*a: object) -> None:
    print("[assemble]", *a, flush=True)


def wc(t: str) -> int:
    t = re.sub(r"<sup>.*?</sup>", "", t)
    t = re.sub(r"^\|.*\|\s*$", "", t, flags=re.M)
    t = re.sub(r"[*_`#>]", "", t)
    return len(re.findall(r"[A-Za-z0-9][A-Za-z0-9\-–—'’.]*", t))


# ------------------------------------------------------------
# 源文稿的公共部分
# ------------------------------------------------------------
def source_parts() -> dict[str, str]:
    text = SRC.read_text(encoding="utf-8")
    text = cit.normalise_a_labels(text)          # [A1] → 真实编号
    heads = list(re.finditer(r"^## (.+?)\s*$", text, re.M))
    parts: dict[str, str] = {"_front": text[: heads[0].start()]}
    for i, m in enumerate(heads):
        end = heads[i + 1].start() if i + 1 < len(heads) else len(text)
        parts[m.group(1).strip()] = text[m.end(): end].strip()
    parts["_full"] = text
    return parts


def title_block(front: str) -> str:
    title = re.search(r"^# (.+?)\s*$", front, re.M).group(1).strip()
    au = re.search(r"^\*\*(.+?)\*\*(.*)$", front, re.M)
    aff = re.search(r"^<sup>1</sup>\s*(.+?)$", front, re.M)
    corr = re.search(r"^Correspondence:\s*(.+?)$", front, re.M)
    return "\n\n".join([
        f"# {title}",
        (au.group(1) + au.group(2)).strip() if au else "",
        f"^1^ {aff.group(1).strip()}" if aff else "",
        f"Correspondence: {corr.group(1).strip()}" if corr else "",
    ])


# ------------------------------------------------------------
# 组装
# ------------------------------------------------------------
def assemble(key: str, parts: dict[str, str]) -> tuple[str, str, dict[str, int]]:
    rev = json.loads((WORK / f"revised_{key}.json").read_text(encoding="utf-8"))
    plan = json.loads((WORK / f"plan_{key}.json").read_text(encoding="utf-8"))
    abstract = plan["abstract"].strip()

    def strip_leading_heading(md: str) -> str:
        """去掉分节 Markdown 自带的一级/二级标题,避免与装配时加的标题重复。

        Some drafts carry their own '## Introduction' heading; the assembler
        adds one, which would otherwise duplicate.
        """
        return re.sub(r"\A\s*#{1,2} [^\n]*\n+", "", md)

    intro, results, disc = (strip_leading_heading(rev[k])
                            for k in ("introduction_md", "results_md", "discussion_md"))
    methods, refs = parts["Methods"], parts["References"]

    is_gcb = key == "GCB"
    if is_gcb:
        ref_map = cit.build_ref_map(parts["_full"])
        conv = lambda t: cit.to_author_date(t, ref_map)   # noqa: E731
        abstract, intro, results, disc, methods = map(conv, (abstract, intro, results, disc, methods))
        # 文献表按姓氏排序、去编号
        entries = []
        for m in re.finditer(r"^(\d+)\.\s+(.+?)(?=\n\d+\.|\n\n|\Z)", refs, re.M | re.S):
            entries.append((ref_map.get(int(m.group(1)), "zzz"), " ".join(m.group(2).split())))
        entries.sort(key=lambda t: t[0].lower())
        refs = "\n\n".join(b for _, b in entries)
        body = [title_block(parts["_front"]), "## Abstract", abstract,
                "## Introduction", intro,
                "## Materials and Methods", methods,
                "## Results", results,
                "## Discussion", disc,
                "## References", refs]
    else:
        body = [title_block(parts["_front"]), "## Abstract", abstract,
                "## Introduction", intro,
                "## Results", results,
                "## Discussion", disc,
                "## References", refs,
                "## Methods", methods]

    back = "\n\n".join(["## Tables", parts.get("Tables", ""),
                        "## Figure legends", parts.get("Figure legends", "")])
    counts = {"abstract": wc(abstract), "introduction": wc(intro),
              "results": wc(results), "discussion": wc(disc), "methods": wc(methods)}
    counts["main_text"] = counts["introduction"] + counts["results"] + counts["discussion"]
    counts["all_inclusive"] = counts["main_text"] + counts["methods"]
    return "\n\n".join(body), back, counts


# ------------------------------------------------------------
# Word 输出(与脚本 200 同一套排版)
# ------------------------------------------------------------
def style_document(path: Path) -> None:
    doc = Document(str(path))
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)

    for st in doc.styles:
        if st.name.startswith("Heading"):
            try:
                st.font.name, st.font.size, st.font.bold = BODY_FONT, Pt(BODY_PT + 1), True
                st.font.color.rgb = RGBColor(0, 0, 0)   # pandoc 默认蓝色标题不适合投稿
            except AttributeError:
                pass

    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1)
        section.top_margin = section.bottom_margin = Inches(1)
        ln = OxmlElement("w:lnNumType")
        for k, v in (("countBy", "1"), ("start", "1"), ("restart", "continuous"), ("distance", "360")):
            ln.set(qn(f"w:{k}"), v)
        section._sectPr.append(ln)
        p = section.footer.paragraphs[0]
        p.alignment = 1
        for kind in ("begin", "instrText", "separate", "end"):
            el = OxmlElement("w:instrText" if kind == "instrText" else "w:fldChar")
            if kind == "instrText":
                el.set(qn("xml:space"), "preserve"); el.text = " PAGE "
            else:
                el.set(qn("w:fldCharType"), kind)
            p.add_run()._r.append(el)

    for tbl in doc.tables:
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4"); e.set(qn("w:color"), "auto")
            borders.append(e)
        tbl._tbl.tblPr.append(borders)
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.size, run.font.name = Pt(9), BODY_FONT

    seen = False
    for para in doc.paragraphs:
        if re.match(r"^Table \d+\.", para.text.strip()):
            if seen:
                para.paragraph_format.page_break_before = True
            seen = True
    doc.save(str(path))


def build(md: str, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    md = cit.__dict__.get("_noop", lambda x: x)(md)
    # <sup>…</sup> → pandoc 原生上标,否则转 docx 时会被整块丢弃
    md = re.sub(r"<sup>(.*?)</sup>",
                lambda m: "^" + m.group(1).strip().replace(" ", r"\ ") + "^", md, flags=re.S)
    tmp = out_path.with_suffix(".md")
    tmp.write_text(md, encoding="utf-8")
    r = subprocess.run(["pandoc", str(tmp), "-o", str(out_path),
                        "--from", "markdown+raw_html+pipe_tables", "--to", "docx", "--standalone"],
                       capture_output=True, text=True)
    if r.returncode != 0:
        sys.exit(f"pandoc 失败 ({out_path.name}): {r.stderr[:400]}")
    style_document(out_path)


def main() -> None:
    parts = source_parts()
    rows = []
    for key, (name, limit) in JOURNALS.items():
        f = WORK / f"revised_{key}.json"
        if not f.exists():
            log(f"跳过 {key}:缺 {f.name}")
            continue
        body_md, back_md, c = assemble(key, parts)
        outdir = SUB / key
        build(body_md, outdir / f"Manuscript_{key}_maintext.docx")
        build(back_md, outdir / f"Manuscript_{key}_tables_and_legends.docx")
        governing = c["all_inclusive"] if key == "GCB" else c["main_text"]
        rows.append([key, name, c["abstract"], c["introduction"], c["results"],
                     c["discussion"], c["methods"], governing, limit,
                     "OK" if governing <= limit else f"超 {governing - limit}"])
        log(f"{key}: 摘要 {c['abstract']} | 正文 {c['main_text']} | 全含 {c['all_inclusive']} | "
            f"约束 {governing}/{limit} {'OK' if governing <= limit else '超限'}")

    with (SUB / "three_versions_wordcount.csv").open("w", newline="", encoding="utf-8") as fh:
        w = csv.writer(fh)
        w.writerow(["key", "journal", "abstract", "introduction", "results",
                    "discussion", "methods", "governing_count", "limit", "status"])
        w.writerows(rows)
    log("完成 ->", SUB)


if __name__ == "__main__":
    main()
